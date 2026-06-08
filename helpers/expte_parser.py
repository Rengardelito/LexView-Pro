# helpers/expte_parser.py
import re
import os
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import io
from datetime import datetime
from docx import Document
import requests
from bs4 import BeautifulSoup
# helpers/expte_parser.py
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import time

# Configuración de Tesseract
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def extraer_texto_docx(ruta_docx):
    """ Lee archivos .docx para extraer texto de párrafos y tablas """
    try:
        doc = Document(ruta_docx)
        texto = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip(): 
                        texto.append(cell.text)
        return '\n'.join(texto).upper()
    except Exception as e:
        print(f"[DOCX] ERROR en {os.path.basename(ruta_docx)}: {e}")
        return ""

def extraer_nro_expte_de_emergencia(ruta_carpeta):
    try:
        archivos = [f for f in os.listdir(ruta_carpeta) if os.path.isfile(os.path.join(ruta_carpeta, f))]
        
        for arch in archivos:
            # Intento 1: Patrón completo NRO-AÑO (ej: 262433-24 o 123456/23)
            match = re.search(r'\b(\d{5,6})\s*[-_/]\s*(\d{2})\b', arch)
            if match:
                return f"{match.group(1)}-{match.group(2)}"
            
            # Intento 2: Solo número largo — SIN inventar el año
            match_simple = re.search(r'\b(\d{5,6})\b', arch)
            if match_simple:
                return match_simple.group(1)  # Devuelve solo el número, ej: "278676"
                
    except Exception as e:
        print(f"[PARSER] Error en escaneo rápido de {ruta_carpeta}: {e}")
    
    return None
def limpiar_secretaria_sucia(txt):
    """ Normaliza nombres de secretarías para la estructura de carpetas """
    if not txt: return "SECRETARIA NRO. 1"
    txt = txt.upper().strip()
    if ' - ' in txt:
        txt = txt.split(' - ')[-1].strip()
    txt = txt.replace("/", "-").replace(":", "-")
    if "SECRETARIA" in txt:
        return txt
    if txt.isdigit():
        return f"SECRETARIA NRO. {txt}"
    return "SECRETARIA NRO. 1"

def extraer_secretaria_pdf(ruta_pdf):
    """ Intenta leer la secretaría desde el contenido del PDF o vía OCR """
    try:
        doc = fitz.open(ruta_pdf)
        texto = ""
        for i in range(min(3, len(doc))):
            texto += doc[i].get_text("text", flags=0) + "\n"
        doc.close()
        
        texto_upper = texto.upper().replace(".", "")

        patrones_sec = [
            r'SECRETAR[IÍ]A\s*N?[°ºRO\.]*\s*(\d+)',
            r'SEC\.?\s*N?[°ºRO\.]*\s*(\d+)',
            r'ANTE\s+M[IÍ]\s*:.*SECRETAR[IÍ]A\s*N?[°ºRO\.]*\s*(\d+)',
            r'JUZGADO.*SECRETAR[IÍ]A\s*N?[°ºRO\.]*\s*(\d+)',
        ]

        for patron in patrones_sec:
            match = re.search(patron, texto_upper, re.IGNORECASE | re.DOTALL)
            if match:
                nro = match.group(1)
                return f"SECRETARIA NRO. {nro}"

        # Si el texto es muy corto o es una imagen, intentar OCR
        if len(texto.strip()) < 50:
            doc = fitz.open(ruta_pdf)
            pix = doc[0].get_pixmap(dpi=200)
            img_data = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_data))
            texto_ocr = pytesseract.image_to_string(img, lang='spa').upper().replace(".", "")
            doc.close()
            for patron in patrones_sec:
                match = re.search(patron, texto_ocr, re.IGNORECASE | re.DOTALL)
                if match:
                    return f"SECRETARIA NRO. {match.group(1)}"

    except Exception as e:
        print(f"[SEC] Error leyendo {ruta_pdf}: {e}")
    return None


# helpers/expte_parser.py



def consultar_forum_corrientes(nro_full):
    """
    Usa tu lógica de Selenium para entrar al Forum y 
    rescatar la ubicación real (Juzgado/Secretaría).
    """
    options = webdriver.ChromeOptions()
    options.add_argument('--headless') # Para que no se abra la ventana cada vez
    driver = webdriver.Chrome(options=options)
    wait = WebDriverWait(driver, 15)
    
    try:
        # 1. Separar nro y año del expediente (ej: 278674-26)
        nro, anio = nro_full.split("-") if "-" in nro_full else (nro_full, "")
        
        # 2. Ir al buscador de causas
        driver.get("https://forumna.juscorrientes.gov.ar/com.forumna.causass")
        
        # 3. Seleccionar Capital (según tu código)
        wait.until(EC.element_to_be_clickable((By.ID, "COMBO_CAUSA_LOCALIDADIDContainer_btnGroupDrop"))).click()
        wait.until(EC.element_to_be_clickable((By.XPATH, "//span[contains(text(), 'Capital')]"))).click()

        # 4. Cargar el número y buscar
        input_nro = wait.until(EC.element_to_be_clickable((By.ID, "vCAUSANRO")))
        input_nro.clear()
        input_nro.send_keys(nro)
        driver.find_element(By.ID, "BTN_SEARCH").click()
        time.sleep(2)

        # 5. Extraer la data de la tabla (Aquí es donde sacamos el JUZGADO)
        # Según tu código, el juzgado suele estar en una de las celdas de la fila
        fila = wait.until(EC.presence_of_element_located((By.XPATH, "//table[@id='Grid1ContainerTbl']//tr[contains(@class, 'Grid')]")))
        
        # Estas posiciones (td[X]) dependen de la tabla real de Forum
        # Normalmente: td[1]=Tipo, td[2]=Nro, td[3]=Año, td[5]=Dependencia/Juzgado
        juzgado_detectado = driver.find_element(By.XPATH, "//table[@id='Grid1ContainerTbl']//tr[2]/td[5]").text.strip()
        
        # Limpiamos el nombre para que sea una carpeta válida
        juzgado_clean = juzgado_detectado.replace("/", "-").upper()

        return {
            "juzgado": juzgado_clean,
            "secretaria": "SECRETARIA UNICA", # O extraer de otra celda
            "caratula": "CARATULA FORUM"
        }

    except Exception as e:
        print(f"⚠️ Error consultando Forum para {nro_full}: {e}")
        return None
    finally:
        driver.quit()