# helpers/cedulas.py
"""
Generador de Cédulas y Mandamientos para LexView Pro.
v2.0 — Múltiples proveídos + copias unificadas en PDF con fojas + márgenes correctos
"""

import os
import re
import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import config


# ── QR ───────────────────────────────────────────────────────────────────────

def _generar_qr_bytes(url: str) -> io.BytesIO | None:
    if not url or not url.strip():
        return None
    try:
        import qrcode
        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_M,
            box_size=10,
            border=4,
        )
        qr.add_data(url.strip())
        qr.make(fit=True)
        img = qr.make_image(fill_color='#1a3a8a', back_color='white')
        buf = io.BytesIO()
        img.save(buf, format='PNG')
        buf.seek(0)
        return buf
    except Exception as e:
        print(f"⚠️ No se pudo generar QR: {e}")
        return None


# ── NUEVO: Unificador de copias en PDF con numeración de fojas ───────────────

def unificar_copias_pdf(rutas_archivos: list[str], ruta_salida: str) -> tuple[str, int]:
    try:
        import fitz

        merger = fitz.open()
        
        for ruta in rutas_archivos:
            if not os.path.exists(ruta):
                print(f"⚠️ No existe: {ruta}")
                continue
            try:
                # repair=True intenta reparar PDFs corruptos
                doc = fitz.open(ruta)
                if doc.is_pdf:
                    doc.save(ruta + ".tmp", garbage=4, deflate=True, clean=True)
                    doc.close()
                    doc = fitz.open(ruta + ".tmp")
                merger.insert_pdf(doc)
                doc.close()
                # Limpiar temp
                tmp = ruta + ".tmp"
                if os.path.exists(tmp):
                    os.remove(tmp)
            except Exception as e:
                print(f"⚠️ No se pudo incluir {os.path.basename(ruta)}: {e}")
                continue

        if len(merger) == 0:
            print("❌ No se pudo unificar ningún PDF")
            merger.close()
            return "", 0

        total_fojas = len(merger)

        for i, page in enumerate(merger):
            texto_foja = f"Foja {i + 1}"
            rect = page.rect
            page.insert_text(
                fitz.Point(rect.width - 80, rect.height - 15),
                texto_foja,
                fontsize=8,
                color=(0.3, 0.3, 0.3),
            )

        os.makedirs(os.path.dirname(ruta_salida), exist_ok=True)
        merger.save(ruta_salida, garbage=4, deflate=True)
        merger.close()

        return ruta_salida, total_fojas

    except Exception as e:
        print(f"❌ Error unificando copias: {e}")
        return "", 0

# ── Helpers de formato ────────────────────────────────────────────────────────

def _set_default_font(doc):
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)


def _parrafo(doc, texto="", bold=False, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=6, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if texto:
        run = p.add_run(texto)
        run.bold      = bold
        run.italic    = italic
        run.font.size = Pt(size)
        run.font.name = 'Arial'
    return p


def _campo(doc, label, valor, size=11):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.size = Pt(size)
    r1.font.name = 'Arial'
    r2 = p.add_run(valor or "")
    r2.font.size = Pt(size)
    r2.font.name = 'Arial'
    return p


def _linea(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '2C3E7A')
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


def _bloque_texto(doc, texto_o_lista, size=11):
    if isinstance(texto_o_lista, str):
        textos = [texto_o_lista] if texto_o_lista.strip() else []
    else:
        textos = [t for t in texto_o_lista if t and t.strip()]

    if not textos:
        return None

    # ── UN SOLO recuadro con todos los proveídos ──────────────────
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(0)
    sep.paragraph_format.space_after  = Pt(2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent       = Cm(0.3)
    p.paragraph_format.right_indent      = Cm(0.3)
    p.paragraph_format.space_before      = Pt(6)
    p.paragraph_format.space_after       = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0)

    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for lado in ['top', 'left', 'bottom', 'right']:
        el = OxmlElement(f'w:{lado}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    '6')
        el.set(qn('w:space'), '6')
        el.set(qn('w:color'), '2C3E7A')
        pBdr.append(el)
    pPr.append(pBdr)

    # Unir todos los textos con doble salto de línea entre ellos
    texto_completo = '\n\n'.join(textos)
    run = p.add_run(texto_completo)
    run.font.size = Pt(size)
    run.font.name = 'Arial'

    sep2 = doc.add_paragraph()
    sep2.paragraph_format.space_before = Pt(2)
    sep2.paragraph_format.space_after  = Pt(6)

    return p


def _bloque_copias(doc, texto_copias, total_fojas=0, size=11):
    """
    Párrafo de copias para traslado.
    Si total_fojas > 0, incluye 'en X fojas' en el texto.
    """
    if total_fojas > 0:
        fojas_txt = f" en {total_fojas} fojas"
        # Insertar "en X fojas" después del último elemento de la lista
        # Formato: "Poder, Titulo, Demanda en X fojas"
        partes = texto_copias.rsplit(',', 1)
        if len(partes) > 1:
            texto_copias = partes[0] + ',' + partes[1].rstrip() + fojas_txt
        else:
            texto_copias = texto_copias.rstrip() + fojas_txt

    texto = (
        f"Además, se acompañan las copias para traslado: {texto_copias}, "
        "las cuales forman parte de esta notificación y sobre las cuales Usted podrá "
        "expresar su aceptación o rechazo a través de un abogado/a de su confianza o "
        "de la defensoría oficial de su localidad siempre que Usted reúna los "
        "requisitos exigidos para su intervención, antes del vencimiento del plazo "
        "ordenado por el Juez/a."
    )
    return _parrafo(doc, texto, size=size, space_after=8,
                    align=WD_ALIGN_PARAGRAPH.JUSTIFY)


def _bloque_acceso(doc, url_drive, qr_bytes, size=11):
    _parrafo(doc, "Para acceder a ellos puede hacerlo a través de:",
             size=size, space_after=4)

    p_a = doc.add_paragraph()
    p_a.paragraph_format.space_after = Pt(2)
    r = p_a.add_run("a) Cargar la URL en su navegador.")
    r.font.size = Pt(size)
    r.font.name = 'Arial'

    p_url = doc.add_paragraph()
    p_url.paragraph_format.space_after = Pt(8)
    r_lbl = p_url.add_run("URL: ")
    r_lbl.bold = True
    r_lbl.font.size = Pt(size)
    r_lbl.font.name = 'Arial'
    r_url = p_url.add_run(url_drive.strip() if url_drive else "___________________________________")
    r_url.font.size = Pt(size)
    r_url.font.name = 'Arial'

    p_b = doc.add_paragraph()
    p_b.paragraph_format.space_after = Pt(6)
    r = p_b.add_run("b) Escanear el siguiente QR:")
    r.font.size = Pt(size)
    r.font.name = 'Arial'

    p_qr = doc.add_paragraph()
    p_qr.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_qr.paragraph_format.space_after = Pt(12)
    run_qr = p_qr.add_run()
    if qr_bytes:
        run_qr.add_picture(qr_bytes, width=Cm(5))
    else:
        run_qr.add_tab()
        run_qr.add_text("[ INSERTAR QR ]")
        run_qr.font.size = Pt(9)


def _pie_cedula(doc, fecha_dia, fecha_mes, fecha_anio):
    _parrafo(doc, "QUEDA UD. LEGALMENTE NOTIFICADO.-",
             bold=True, size=11, space_before=6, space_after=4,
             align=WD_ALIGN_PARAGRAPH.LEFT)
    fecha = f"Corrientes, {fecha_dia} de {fecha_mes} {fecha_anio}.-"
    _parrafo(doc, fecha, size=11, space_after=30, align=WD_ALIGN_PARAGRAPH.LEFT)


def _pie_mandamiento(doc, fecha_dia, fecha_mes, fecha_anio, localidad="Corrientes"):
    fecha = (
        f"Dado, sellado y firmado en la Sala de mi Público Despacho, "
        f"a los {fecha_dia} días del mes de {fecha_mes} del año {fecha_anio}.-"
    )
    _parrafo(doc, fecha, size=11, space_before=10, space_after=30,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY)


# ── Función principal ─────────────────────────────────────────────────────────

def generar_cedula(datos: dict, tipo: str, ruta_salida: str) -> str:
    """
    Genera el .docx y lo guarda en ruta_salida.

    NUEVO en datos:
        "textos_providencia": lista de strings (múltiples proveídos)
                              Si se pasa, tiene precedencia sobre "texto_providencia"
        "rutas_copias":       lista de rutas de PDFs a unificar como copias
                              Si se pasa, genera el PDF unificado automáticamente
        "total_fojas":        se calcula automáticamente si se pasa rutas_copias
    """
    doc = Document()
    _set_default_font(doc)

    # ── Márgenes: Sup 4.5 | Inf 2.5 | Der 1.5 | Izq 4.5 (espejados/simétricos)
    for section in doc.sections:
        section.page_width    = Cm(21)
        section.page_height   = Cm(29.7)
        section.top_margin    = Cm(4.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(4.5)
        section.right_margin  = Cm(1.5)

    # ── Normalizar proveídos: acepta string o lista ──────────────────────────
    if "textos_providencia" in datos and datos["textos_providencia"]:
        textos_prov = datos["textos_providencia"]
        if isinstance(textos_prov, str):
            textos_prov = [textos_prov]
    else:
        texto_unico = datos.get("texto_providencia", "")
        textos_prov = [texto_unico] if texto_unico else []

    datos["_textos_prov_norm"] = textos_prov

    # ── Unificar copias si vienen rutas ─────────────────────────────────────
    total_fojas = datos.get("total_fojas", 0)
    rutas_copias = datos.get("rutas_copias", [])
    print(f"[DEBUG unificar] rutas_copias en generar_cedula: {rutas_copias}")
    if rutas_copias and not total_fojas:
        ruta_pdf_copias = ruta_salida.replace(".docx", "_copias.pdf")
        print(f"[DEBUG unificar] guardando PDF copias en: {ruta_pdf_copias}")
        _, total_fojas = unificar_copias_pdf(rutas_copias, ruta_pdf_copias)
        print(f"[DEBUG unificar] total_fojas: {total_fojas}")
        datos["_ruta_pdf_copias"] = ruta_pdf_copias

    datos["_total_fojas"] = total_fojas

    # ── Generar QR ───────────────────────────────────────────────────────────
    url_drive = datos.get("url_drive", "").strip()
    qr_bytes  = _generar_qr_bytes(url_drive) if url_drive else None
    hay_copias = bool(datos.get("copias_traslado", "").strip()) or bool(rutas_copias)

    if tipo == "cedula_local":
        _cedula_local(doc, datos, url_drive, qr_bytes, hay_copias)
    elif tipo == "cedula_ley":
        _cedula_ley(doc, datos, url_drive, qr_bytes, hay_copias)
    elif tipo == "mandamiento_local":
        _mandamiento_local(doc, datos, url_drive, qr_bytes, hay_copias)
    elif tipo == "mandamiento_ley":
        _mandamiento_ley(doc, datos, url_drive, qr_bytes, hay_copias)
    else:
        raise ValueError(f"Tipo desconocido: {tipo}")

    doc.save(ruta_salida)
    return ruta_salida


# ── Cédula Local ──────────────────────────────────────────────────────────────

def _cedula_local(doc, d, url_drive, qr_bytes, hay_copias):
    _campo(doc, "Juzgado",                   d.get("juzgado", ""))
    _campo(doc, "N° Expte.",                 d.get("nro_expte", ""))
    _campo(doc, "Carátula",                  d.get("caratula", ""))
    _campo(doc, "Juez/a Dr./Dra.",           d.get("juez", ""))
    _campo(doc, "Secretaría del Dr./Dra.",   d.get("secretaria", ""))
    _campo(doc, "Domicilio del Juzgado",     d.get("domicilio_juzgado", ""))
    _campo(doc, "Objeto de la Notificación", d.get("objeto_notificacion", ""))
    _linea(doc)
    _parrafo(doc, "CÉDULA", bold=True, size=14,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=6)
    _campo(doc, "DESTINATARIO/S",         d.get("destinatario", ""))
    _campo(doc, "DOMICILIO",              d.get("domicilio", ""))
    _campo(doc, "CARÁCTER DEL DOMICILIO", d.get("caracter_domicilio", ""))
    _campo(doc, "LOCALIDAD",              d.get("localidad", ""))
    _parrafo(doc, space_after=4)
    _parrafo(doc,
        "Por la presente cédula se pone en su conocimiento que en el proceso judicial "
        "mencionado se ordenó notificarle por este medio lo siguiente:",
        size=11, space_after=6)
    _bloque_texto(doc, d.get("_textos_prov_norm", [d.get("texto_providencia", "")]))
    if hay_copias:
        _bloque_copias(doc, d.get("copias_traslado", ""), total_fojas=d.get("_total_fojas", 0))
    if hay_copias or url_drive:
        _bloque_acceso(doc, url_drive, qr_bytes)
    _pie_cedula(doc, d.get("fecha_dia",""), d.get("fecha_mes",""), d.get("fecha_anio","2026"))


# ── Cédula Ley 22172/3556 ────────────────────────────────────────────────────

def _cedula_ley(doc, d, url_drive, qr_bytes, hay_copias):
    _campo(doc, "Organismo emisor",        d.get("juzgado", ""))
    _campo(doc, "Correo institucional",    d.get("correo_juzgado", ""))
    _campo(doc, "Domicilio y N° de teléfono del Juzgado",
           f"{d.get('domicilio_juzgado','')}  {d.get('tel_juzgado','')}".strip())
    _campo(doc, "Dependencia destinataria", d.get("dependencia_dest", ""))
    _campo(doc, "N° Expte.",               d.get("nro_expte", ""))
    _campo(doc, "Carátula",                d.get("caratula", ""))
    _campo(doc, "Personas Autorizadas",    d.get("personas_autorizadas", ""))
    _campo(doc, "Objeto de la Notificación", d.get("objeto_notificacion", ""))
    _linea(doc)
    _parrafo(doc, "CÉDULA LEY 22172/3556", bold=True, size=14,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=6)
    _campo(doc, "DESTINATARIO/S",         d.get("destinatario", ""))
    _campo(doc, "DOMICILIO",              d.get("domicilio", ""))
    _campo(doc, "CARÁCTER DEL DOMICILIO", d.get("caracter_domicilio", ""))
    _campo(doc, "LOCALIDAD",              d.get("localidad", ""))
    _parrafo(doc, space_after=4)
    _parrafo(doc,
        "Por la presente cédula se pone en su conocimiento que en el proceso judicial "
        "mencionado se ordenó notificarle por este medio lo siguiente:",
        size=11, space_after=6)
    _bloque_texto(doc, d.get("_textos_prov_norm", [d.get("texto_providencia", "")]))
    if hay_copias:
        _bloque_copias(doc, d.get("copias_traslado", ""), total_fojas=d.get("_total_fojas", 0))
    if hay_copias or url_drive:
        _bloque_acceso(doc, url_drive, qr_bytes)
    _pie_cedula(doc, d.get("fecha_dia",""), d.get("fecha_mes",""), d.get("fecha_anio","2026"))


# ── Mandamiento Local ─────────────────────────────────────────────────────────

def _mandamiento_local(doc, d, url_drive, qr_bytes, hay_copias):
    _campo(doc, "Juzgado",                  d.get("juzgado", ""))
    _campo(doc, "N° Expte.",                d.get("nro_expte", ""))
    _campo(doc, "Carátula",                 d.get("caratula", ""))
    _campo(doc, "Juez/a Dr./Dra.",          d.get("juez", ""))
    _campo(doc, "Secretaría del Dr./Dra.",  d.get("secretaria", ""))
    _campo(doc, "Domicilio del Juzgado",    d.get("domicilio_juzgado", ""))
    _linea(doc)
    _parrafo(doc, "MODELO DE MANDAMIENTO ESTANDARIZADO", bold=True, size=14,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=10)
    domicilio = d.get("domicilio", "……………………………………")
    _parrafo(doc,
        f"El Sr. Oficial de Justicia se hará presente en el domicilio ubicado en "
        f"{domicilio}, y PROCEDERÁ a dar cumplimiento de lo ordenado en el expediente "
        f"mencionado mediante la Resolución N° ……………, que a continuación se transcribe:",
        size=11, space_after=8)
    _bloque_texto(doc, d.get("_textos_prov_norm", [d.get("texto_providencia", "")]))
    if hay_copias:
        _bloque_copias(doc, d.get("copias_traslado", ""), total_fojas=d.get("_total_fojas", 0),
                       size=11)
    if hay_copias or url_drive:
        _bloque_acceso(doc, url_drive, qr_bytes)
    _pie_mandamiento(doc, d.get("fecha_dia",""), d.get("fecha_mes",""), d.get("fecha_anio","2026"))


# ── Mandamiento Ley 22172/3556 ────────────────────────────────────────────────

def _mandamiento_ley(doc, d, url_drive, qr_bytes, hay_copias):
    _campo(doc, "Organismo emisor",        d.get("juzgado", ""))
    _campo(doc, "Correo institucional",    d.get("correo_juzgado", ""))
    _campo(doc, "Domicilio y N° de tel. del Juzgado",
           f"{d.get('domicilio_juzgado','')}  {d.get('tel_juzgado','')}".strip())
    _campo(doc, "Dependencia destinataria", d.get("dependencia_dest", ""))
    _campo(doc, "N° Expte.",               d.get("nro_expte", ""))
    _campo(doc, "Carátula",                d.get("caratula", ""))
    _linea(doc)
    _parrafo(doc, "MANDAMIENTO LEY (22172/3556)", bold=True, size=14,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=10)
    domicilio = d.get("domicilio", "……………………………………")
    _parrafo(doc,
        f"El Sr. Oficial de Justicia se hará presente en el domicilio ubicado en "
        f"{domicilio}, y PROCEDERÁ a dar cumplimiento de lo ordenado en el expediente "
        f"mencionado mediante la Resolución N° ……………, que a continuación se transcribe:",
        size=11, space_after=8)
    _bloque_texto(doc, d.get("_textos_prov_norm", [d.get("texto_providencia", "")]))
    if hay_copias:
        _bloque_copias(doc, d.get("copias_traslado", ""), total_fojas=d.get("_total_fojas", 0),
                       size=11)
    if hay_copias or url_drive:
        _bloque_acceso(doc, url_drive, qr_bytes)
    _pie_mandamiento(doc, d.get("fecha_dia",""), d.get("fecha_mes",""), d.get("fecha_anio","2026"),
                     localidad=d.get("localidad","Corrientes"))


# ── Extractor y listador de proveídos ────────────────────────────────────────

def extraer_texto_proveido(ruta_pdf: str) -> str:
    try:
        import fitz
        doc = fitz.open(ruta_pdf)
        texto = ""
        for page in doc:
            texto += page.get_text()
        doc.close()
        texto = re.sub(r'\n{3,}', '\n\n', texto.strip())
        return texto
    except Exception as e:
        print(f"⚠️ No se pudo extraer texto de {ruta_pdf}: {e}")
        return ""


def listar_proveidos(ruta_carpeta: str) -> list:
    if not os.path.isdir(ruta_carpeta):
        return []

    patron_fecha_escrito = re.compile(r'^\d{1,2}[/_]\d{1,2}[/_]\d{2,4}')
    proveidos = []

    for nombre in sorted(os.listdir(ruta_carpeta), reverse=True):
        if not nombre.lower().endswith('.pdf'):
            continue
        if nombre == 'caratula_pro.pdf':
            continue

        partes = nombre.replace('.pdf', '').split(' - ', 1)
        if len(partes) < 2:
            continue

        fecha_str = partes[0].strip()
        extracto  = partes[1].strip()

        if patron_fecha_escrito.match(extracto):
            continue

        try:
            fecha_display = datetime.strptime(fecha_str, "%Y-%m-%d").strftime("%d/%m/%Y")
        except Exception:
            fecha_display = fecha_str

        proveidos.append({
            "nombre":   nombre,
            "ruta":     os.path.join(ruta_carpeta, nombre),
            "fecha":    fecha_display,
            "extracto": extracto
        })

    return proveidos